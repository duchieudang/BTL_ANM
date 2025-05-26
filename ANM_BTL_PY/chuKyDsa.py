import hashlib
import random
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from docx import Document
import fitz  # PyMuPDF
# --- DSA functions ---

def hash_message(message):
    sha256 = hashlib.sha256()
    sha256.update(message.encode('utf-8'))
    return int(sha256.hexdigest(), 16)


def hash_signature(r, s):
    sha1 = hashlib.sha1()
    sha1.update((str(r) + str(s)).encode('utf-8'))
    return sha1.hexdigest()

def calculate_g(p, q, h):
    return pow(h, (p - 1) // q, p)

def generate_k(q):
    while True:
        k = random.randint(1, q - 1)
        if 1 <= k < q:
            return k

def modinv(a, m):
    g, x, _ = extended_gcd(a, m)
    if g != 1:
        raise Exception("Không tìm được nghịch đảo modular")
    return x % m

def extended_gcd(a, b):
    if a == 0:
        return (b, 0, 1)
    else:
        g, y, x = extended_gcd(b % a, a)
        return (g, x - (b // a) * y, y)

def sign(message, p, q, g, x, k):
    r = pow(g, k, p) % q
    if r == 0:
        raise ValueError("r = 0, chọn k khác.")
    h = hash_message(message)
    k_inv = modinv(k, q)
    s = (k_inv * (h + x * r)) % q
    if s == 0:
        raise ValueError("s = 0, chọn k khác.")
    return (r, s)

def verify(message, p, q, g, y, r, s):
    if not (0 < r < q) or not (0 < s < q):
        return False
    w = modinv(s, q)
    h = hash_message(message)
    u1 = (h * w) % q
    u2 = (r * w) % q
    v = ((pow(g, u1, p) * pow(y, u2, p)) % p) % q
    return v == r

# --- GUI ---

class DSAGui:
    def __init__(self, root):
        self.root = root
        root.title("DSA Digital Signature Tool")
        root.geometry("900x750")
        root.configure(bg="#f9f9f9")

        self.setup_style()

        self.r = None
        self.s = None
        self.signature_hash = None
        self.message = None

        self.p_var = tk.StringVar()
        self.q_var = tk.StringVar()
        self.h_var = tk.StringVar()
        self.g_var = tk.StringVar()
        self.x_var = tk.StringVar()
        self.y_var = tk.StringVar()

        self.setup_widgets()

    def setup_style(self):
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("TButton", padding=6, font=('Segoe UI', 10))
        style.configure("TLabel", font=('Segoe UI', 10))
        style.configure("TEntry", font=('Segoe UI', 10))

    def setup_widgets(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(expand=True, fill='both')

        self.sign_tab = ttk.Frame(notebook)
        self.verify_tab = ttk.Frame(notebook)
        notebook.add(self.sign_tab, text="🔐 Ký văn bản")
        notebook.add(self.verify_tab, text="✅ Xác thực chữ ký")

        self.create_param_section(self.sign_tab)
        self.create_sign_section(self.sign_tab)
        self.create_verify_section(self.verify_tab)

    def create_param_section(self, parent):
        frame = ttk.LabelFrame(parent, text="Tham số DSA", padding=10)
        frame.pack(fill="x", padx=10, pady=10)

        fields = [
            ("p:", self.p_var),
            ("q:", self.q_var),
            ("h:", self.h_var),
            ("g:", self.g_var),
            ("x (private key):", self.x_var),
            ("y (public key):", self.y_var)
        ]

        for i, (label, var) in enumerate(fields):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky="e", padx=5, pady=2)
            ttk.Entry(frame, textvariable=var, width=80).grid(row=i, column=1, sticky="w", padx=5, pady=2)

        ttk.Button(frame, text="🎲 Sinh ngẫu nhiên tham số", command=self.generate_all_params).grid(row=6, column=0, columnspan=2, pady=5)

    def load_pdf_file(self, text_widget):
        file_path = filedialog.askopenfilename(
            title="Chọn file PDF (.pdf)",
            filetypes=[("PDF Files", "*.pdf")]
        )
        if not file_path:
            return
        try:
            doc = fitz.open(file_path)
            full_text = []
            for page in doc:
                full_text.append(page.get_text())
            content = '\n'.join(full_text)
            text_widget.delete("1.0", tk.END)
            text_widget.insert(tk.END, content)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file PDF.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file PDF: {e}")


    def load_excel_file(self, text_widget):
        from tkinter import filedialog, messagebox
        file_path = filedialog.askopenfilename(
            title="Chọn file Excel (.xlsx)",
            filetypes=[("Excel Files", "*.xlsx")]
        )
        if not file_path:
            return
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active
            content = []
            for row in sheet.iter_rows(values_only=True):
                content.append('\t'.join([str(cell) if cell is not None else '' for cell in row]))
            text_data = '\n'.join(content)
            text_widget.delete("1.0", "end")
            text_widget.insert("end", text_data)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file Excel.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Excel: {e}")

    def create_sign_section(self, parent):
        frame = ttk.LabelFrame(parent, text="Ký nội dung", padding=10)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ttk.Label(frame, text="Nội dung cần ký:").pack(anchor='w')

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(fill='x', pady=(0, 5))

        ttk.Button(btn_frame, text="📂 Mở file Word (.docx)", command=self.load_word_file_to_sign).pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📂 Mở file PDF (.pdf)", command=lambda: self.load_pdf_file(self.text_input)).pack(side='left', padx=(0, 5))
        ttk.Button(btn_frame, text="📂 Mở file Excel (.xlsx)", command=lambda: self.load_excel_file(self.text_input)).pack(side='left', padx=(0, 5))

        ttk.Button(btn_frame, text="👁️ Xem chi tiết", command=lambda: self.show_detail_popup(self.text_input.get("1.0", tk.END))).pack(side='left')

        self.text_input = tk.Text(frame, height=6, wrap='word')
        self.text_input.pack(fill="x", pady=5)

        ttk.Button(frame, text="🖊️ Tạo chữ ký", command=self.sign_message).pack(pady=5)

        ttk.Label(frame, text="Kết quả (r, s, hash):").pack(anchor='w', pady=(10, 0))
        self.signature_text = tk.Text(frame, height=5, wrap='word', state='disabled')
        self.signature_text.pack(fill="x", pady=5)

    def show_detail_popup(self, content):
        detail_win = tk.Toplevel(self.root)
        detail_win.title("Chi tiết nội dung")
        detail_win.geometry("600x400")
        text = tk.Text(detail_win, wrap='word')
        text.insert(tk.END, content)
        text.config(state='disabled')
        text.pack(fill='both', expand=True, padx=10, pady=10)
    def create_verify_section(self, parent):
        frame = ttk.LabelFrame(parent, text="Xác thực chữ ký", padding=10)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        ttk.Label(frame, text="Nội dung cần xác thực:").pack(anchor='w')

        self.verify_text_input = tk.Text(frame, height=8, wrap='word')
        self.verify_text_input.pack(fill="x", pady=5)

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(pady=5)
        ttk.Button(btn_frame, text="📂 Mở file Word (.docx)", command=self.load_word_file).pack(side='left', padx=(0,5))
        ttk.Button(btn_frame, text="📂 Mở file PDF (.pdf)", command=lambda: self.load_pdf_file(self.verify_text_input)).pack(side='left', padx=(0,5))
        ttk.Button(btn_frame, text="📂 Mở file Excel (.xlsx)", command=lambda: self.load_excel_file(self.text_input)).pack(side='left', padx=(0, 5))        
        ttk.Button(btn_frame, text="👁️ Xem chi tiết", command=lambda: self.show_detail_popup(self.verify_text_input.get("1.0", tk.END))).pack(side='left')

        ttk.Label(frame, text="Hash chữ ký:").pack(anchor='w')
        self.verify_hash_entry = ttk.Entry(frame, width=80)
        self.verify_hash_entry.pack(pady=5)

        ttk.Button(frame, text="✅ Xác thực", command=self.verify_signature).pack(pady=5)
        self.result_label = tk.Label(frame, text="", font=('Segoe UI', 11, 'bold'), bg="#f9f9f9")
        self.result_label.pack()

    def generate_all_params(self):
        self.q = random.choice([101, 103, 107])
        self.p = self.q * random.randint(70, 100) + 1
        while not self.is_prime(self.p):
            self.p += self.q
        self.h = random.randint(2, self.p - 2)
        self.g = calculate_g(self.p, self.q, self.h)
        self.x = random.randint(1, self.q - 1)
        self.y = pow(self.g, self.x, self.p)

        self.p_var.set(str(self.p))
        self.q_var.set(str(self.q))
        self.h_var.set(str(self.h))
        self.g_var.set(str(self.g))
        self.x_var.set(str(self.x))
        self.y_var.set(str(self.y))

    def is_prime(self, n):
        if n < 2: return False
        for i in range(2, int(n ** 0.5) + 1):
            if n % i == 0: return False
        return True

    def load_word_file_to_sign(self):
        file_path = filedialog.askopenfilename(
            title="Chọn file Word (.docx)",
            filetypes=[("Word Documents", "*.docx")]
        )
        if not file_path:
            return
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            self.text_input.delete("1.0", tk.END)
            self.text_input.insert(tk.END, content)
            messagebox.showinfo("Thành công", "Đã tải nội dung từ file Word vào phần ký.")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Word: {e}")

    def load_word_file(self):
        file_path = filedialog.askopenfilename(
            title="Chọn file Word (.docx)",
            filetypes=[("Word Documents", "*.docx")]
        )
        if not file_path:
            return
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = '\n'.join(full_text)
            self.verify_text_input.delete("1.0", tk.END)
            self.verify_text_input.insert(tk.END, content)
            self.result_label.config(text="Đã tải nội dung từ file Word.", fg="blue")
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể đọc file Word: {e}")

    def sign_message(self):
        try:
            message = self.text_input.get("1.0", tk.END).strip()
            if not message:
                messagebox.showwarning("Cảnh báo", "Vui lòng nhập nội dung!")
                return

            self.p = int(self.p_var.get())
            self.q = int(self.q_var.get())
            self.h = int(self.h_var.get())
            self.g = int(self.g_var.get())
            self.x = int(self.x_var.get())
            self.y = int(self.y_var.get())

            self.k = generate_k(self.q)
            self.r, self.s = sign(message, self.p, self.q, self.g, self.x, self.k)
            self.signature_hash = hash_signature(self.r, self.s)

            self.signature_text.config(state='normal')
            self.signature_text.delete("1.0", tk.END)
            self.signature_text.insert(tk.END, f"r = {self.r}\ns = {self.s}\nHash: {self.signature_hash}")
            self.signature_text.config(state='disabled')

            self.verify_text_input.delete("1.0", tk.END)
            self.verify_text_input.insert(tk.END, message)
            self.verify_hash_entry.delete(0, tk.END)
            self.verify_hash_entry.insert(0, self.signature_hash)

            self.message = message
            self.result_label.config(text="", fg="blue")
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

    def verify_signature(self):
        try:
            message = self.verify_text_input.get("1.0", tk.END).strip()
            signature_hash_input = self.verify_hash_entry.get().strip()

            if not message or not signature_hash_input:
                messagebox.showwarning("Cảnh báo", "Vui lòng nhập đủ nội dung và hash chữ ký!")
                return

            self.p = int(self.p_var.get())
            self.q = int(self.q_var.get())
            self.h = int(self.h_var.get())
            self.g = int(self.g_var.get())
            self.x = int(self.x_var.get())
            self.y = int(self.y_var.get())

            # Nếu chưa ký lần nào hoặc chưa có r,s
            if self.r is None or self.s is None:
                messagebox.showwarning("Cảnh báo", "Chưa có chữ ký để xác thực! Vui lòng ký trước.")
                return

            # Kiểm tra hash chữ ký
            if signature_hash_input != self.signature_hash:
                self.result_label.config(text="Hash chữ ký không đúng!", fg="red")
                return

            verified = verify(message, self.p, self.q, self.g, self.y, self.r, self.s)
            if verified:
                self.result_label.config(text="Xác thực chữ ký thành công ✅", fg="green")
            else:
                self.result_label.config(text="Xác thực chữ ký thất bại ❌", fg="red")
        except Exception as e:
            messagebox.showerror("Lỗi", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = DSAGui(root)
    root.mainloop()
#pip install python-docx
#pip install pymupdf
